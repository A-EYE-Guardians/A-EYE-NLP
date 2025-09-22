# loader/processing_document.py
"""
웹 전용 문서 처리 파이프라인

- 입력: URL (문서/웹페이지)
- 처리:
  1) URL -> 임시 파일 다운로드
  2) 확장자/Content-Type 기반으로 적합한 extractor로 텍스트 추출
  3) 텍스트 클린업, 청크 분할
  4) (옵션) 임베딩 생성
- 반환: dict {
    "source": url,
    "raw_text": "...",
    "chunks": [{"id":0,"text":..., "start":, "end":}, ...],
    "embeddings": [np.ndarray] or None,
    "status": "ok" / "error",
    "error_msg": optional
  }
"""

import os
import re
import tempfile
import shutil
import logging
import requests
from pathlib import Path
from urllib.parse import urlparse
from typing import List, Dict, Optional, Tuple, Any
from dotenv import load_dotenv
import numpy as np
import time

# text/html parsing
from bs4 import BeautifulSoup

# PDF parsers
try:
    import pdfplumber
except Exception:
    pdfplumber = None
try:
    import pypdf
except Exception:
    pypdf = None

# docx
try:
    import docx
except Exception:
    docx = None

# hwp
try:
    import pyhwp
except Exception:
    pyhwp = None

# OCR fallback (pdf -> images -> tesseract)
try:
    from pdf2image import convert_from_path
    import pytesseract
except Exception:
    convert_from_path = None
    pytesseract = None

# optional mammoth for docx->html
try:
    import mammoth
except Exception:
    mammoth = None

# embeddings (OpenAI)
from openai import OpenAI
load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
_client = OpenAI(api_key=OPENAI_API_KEY) if OPENAI_API_KEY else None

LOGGER = logging.getLogger(__name__)
LOGGER.setLevel(logging.INFO)

# directory to cache raw texts (optional)
SAVED_DOCS_DIR = Path("./saved_documents")
SAVED_DOCS_DIR.mkdir(exist_ok=True, parents=True)

# -------------------------
# Helpers
# -------------------------
def is_url(text: str) -> bool:
    try:
        p = urlparse(text)
        return bool(p.scheme and p.netloc)
    except Exception:
        return False

def _download_to_temp(url: str, timeout: int = 15) -> Tuple[Path, str]:
    """
    Download URL to a temporary file. Return (Path, content_type)
    """
    tmp = tempfile.NamedTemporaryFile(delete=False)
    tmp_path = Path(tmp.name)
    headers = {"User-Agent": "Mozilla/5.0 (compatible; A-EYE/1.0)"}
    try:
        r = requests.get(url, headers=headers, timeout=timeout, stream=True)
        r.raise_for_status()
        # guess ext from content-type or url
        content_type = r.headers.get("content-type", "")
        ext = ""
        if "pdf" in content_type or url.lower().endswith(".pdf"):
            ext = ".pdf"
        elif "html" in content_type or url.lower().endswith((".html", ".htm")):
            ext = ".html"
        elif "msword" in content_type or url.lower().endswith(".docx"):
            ext = ".docx"
        elif url.lower().endswith(".txt") or "text/plain" in content_type:
            ext = ".txt"
        elif url.lower().endswith((".hwp", ".hwpx")) or "hwp" in content_type:
            ext = ".hwp"
        else:
            # fallback to html
            ext = ".html"
        tmp.close()
        final_tmp = tmp_path.with_suffix(ext)
        with open(final_tmp, "wb") as f:
            for chunk in r.iter_content(chunk_size=8192):
                if chunk:
                    f.write(chunk)
        return final_tmp, content_type
    except Exception as e:
        # ensure cleanup
        try:
            tmp.close()
            tmp_path.unlink(missing_ok=True)
        except Exception:
            pass
        raise

def _clean_text(text: str) -> str:
    if not text:
        return ""
    # normalize newlines and whitespace
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r'\n\s*\n+', '\n\n', text)
    text = re.sub(r'[ \t]{2,}', ' ', text)
    return text.strip()

# -------------------------
# Extractors
# -------------------------
def _extract_text_from_html_file(path: Path) -> str:
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            html = f.read()
    except Exception:
        # binary read fallback
        with open(path, "rb") as f:
            html = f.read().decode(errors="ignore")
    soup = BeautifulSoup(html, "html.parser")
    for s in soup(["script", "style", "header", "footer", "nav", "aside"]):
        s.decompose()
    text = soup.get_text(separator="\n")
    return _clean_text(text)

def _extract_text_from_pdf_file(path: Path) -> str:
    texts = []
    if pdfplumber:
        try:
            with pdfplumber.open(path) as pdf:
                for p in pdf.pages:
                    texts.append(p.extract_text() or "")
            joined = "\n".join(texts).strip()
            if joined:
                return _clean_text(joined)
        except Exception as e:
            LOGGER.debug("pdfplumber failed: %s", e)
    if pypdf:
        try:
            reader = pypdf.PdfReader(str(path))
            for page in reader.pages:
                try:
                    txt = page.extract_text() or ""
                except Exception:
                    txt = ""
                texts.append(txt)
            joined = "\n".join(texts).strip()
            if joined:
                return _clean_text(joined)
        except Exception as e:
            LOGGER.debug("pypdf failed: %s", e)
    # OCR fallback if available
    if convert_from_path and pytesseract:
        try:
            pages = convert_from_path(str(path), dpi=200)
            ocr_texts = []
            for img in pages:
                ocr_texts.append(pytesseract.image_to_string(img, lang="kor+eng"))
            joined = "\n".join(ocr_texts).strip()
            return _clean_text(joined)
        except Exception as e:
            LOGGER.debug("PDF OCR failed: %s", e)
    return ""

def _extract_text_from_docx_file(path: Path) -> str:
    # try python-docx first
    if docx:
        try:
            d = docx.Document(str(path))
            paragraphs = [p.text for p in d.paragraphs if p.text.strip()]
            joined = "\n".join(paragraphs).strip()
            if joined:
                return _clean_text(joined)
        except Exception as e:
            LOGGER.debug("python-docx failed: %s", e)
    # try mammoth
    if mammoth:
        try:
            with open(path, "rb") as f:
                res = mammoth.extract_raw_text(f)
                txt = res.value or ""
                return _clean_text(txt)
        except Exception as e:
            LOGGER.debug("mammoth failed: %s", e)
    # fallback binary decode
    try:
        with open(path, "rb") as f:
            raw = f.read().decode(errors="ignore")
            return _clean_text(raw)
    except Exception:
        return ""

def _extract_text_from_txt_file(path: Path) -> str:
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return _clean_text(f.read())
    except Exception:
        try:
            with open(path, "rb") as f:
                return _clean_text(f.read().decode(errors="ignore"))
        except Exception:
            return ""

def _extract_text_from_hwp_file(path: Path) -> str:
    # pyhwp usage is environment-specific; try basic approach
    if pyhwp:
        try:
            doc = pyhwp.HWPDocument(str(path))
            # pyhwp API complex — fallback to empty unless specific environment
            return ""
        except Exception as e:
            LOGGER.debug("pyhwp failed: %s", e)
    # try libreoffice conversion if available
    try:
        outdir = tempfile.mkdtemp()
        cmd = f'soffice --headless --convert-to docx "{path}" --outdir "{outdir}"'
        os.system(cmd)
        for f in os.listdir(outdir):
            if f.endswith(".docx"):
                return _extract_text_from_docx_file(Path(outdir) / f)
    except Exception as e:
        LOGGER.debug("libreoffice conversion failed: %s", e)
    return ""

# -------------------------
# Chunking + Embeddings
# -------------------------
def chunk_text(text: str, chunk_size_chars: int = 2000, overlap: int = 200) -> List[Dict]:
    if not text:
        return []
    chunks = []
    start = 0
    text_len = len(text)
    idx = 0
    while start < text_len:
        end = min(start + chunk_size_chars, text_len)
        ctext = text[start:end]
        chunks.append({"id": idx, "text": ctext, "start": start, "end": end})
        idx += 1
        start = end - overlap
        if start < 0:
            start = 0
    return chunks

def embed_texts(texts: List[str], batch_size: int = 8, model: str = "text-embedding-3-small") -> Optional[List[np.ndarray]]:
    if _client is None:
        LOGGER.info("OpenAI client not configured; skipping embeddings.")
        return None
    embeddings = []
    for i in range(0, len(texts), batch_size):
        batch = texts[i:i+batch_size]
        try:
            resp = _client.embeddings.create(model=model, input=batch)
            for item in resp.data:
                embeddings.append(np.array(item.embedding, dtype="float32"))
        except Exception as e:
            LOGGER.warning("Embedding request failed: %s", e)
            # align with zeros fallback
            dim = 1536
            for _ in batch:
                embeddings.append(np.zeros((dim,), dtype="float32"))
    return embeddings

# -------------------------
# Top-level single document processing (URL)
# -------------------------
def process_document_for_llm(source_url: str, from_url: bool = True,
                             do_embed: bool = False,
                             chunk_size_chars: int = 2000,
                             overlap: int = 200,
                             timeout: int = 15) -> Dict[str, Any]:
    """
    Process a single URL -> return dict.
    """
    result = {
        "source": source_url,
        "raw_text": "",
        "chunks": [],
        "embeddings": None,
        "status": "error",
        "error_msg": None
    }

    if not source_url or not is_url(source_url):
        result["error_msg"] = "source is not a valid URL"
        return result

    tmp_path = None
    try:
        tmp_path, content_type = _download_to_temp(source_url, timeout=timeout)
        ext = tmp_path.suffix.lower()
        text = ""
        if ext in (".html", ".htm"):
            text = _extract_text_from_html_file(tmp_path)
        elif ext == ".pdf":
            text = _extract_text_from_pdf_file(tmp_path)
        elif ext in (".docx",):
            text = _extract_text_from_docx_file(tmp_path)
        elif ext in (".txt", ".md"):
            text = _extract_text_from_txt_file(tmp_path)
        elif ext in (".hwp", ".hwpx"):
            text = _extract_text_from_hwp_file(tmp_path)
        else:
            # try html then pdf
            text = _extract_text_from_html_file(tmp_path)
            if not text.strip():
                text = _extract_text_from_pdf_file(tmp_path)

        text = _clean_text(text)
        result["raw_text"] = text
        result["chunks"] = chunk_text(text, chunk_size_chars=chunk_size_chars, overlap=overlap)

        if do_embed and result["chunks"] and _client is not None:
            texts = [c["text"] for c in result["chunks"]]
            embs = embed_texts(texts)
            result["embeddings"] = embs
            # attach embedding refs to chunks optionally
            if embs:
                for c, e in zip(result["chunks"], embs):
                    c["embedding"] = e.tolist() if isinstance(e, np.ndarray) else None

        result["status"] = "ok"
        return result

    except Exception as e:
        LOGGER.exception("process_document_for_llm failed for %s: %s", source_url, e)
        result["error_msg"] = str(e)
        result["status"] = "error"
        return result

    finally:
        # cleanup temp file
        try:
            if tmp_path and tmp_path.exists():
                tmp_path.unlink(missing_ok=True)
        except Exception:
            pass

# -------------------------
# Multi-document
# -------------------------
def process_multiple_documents(
    sources: List[str],
    from_url: bool = True,
    do_embed: bool = False,
    chunk_size_chars: int = 2000,
    overlap: int = 200
) -> List[Dict[str, Any]]:
    results = []
    for src in sources:
        try:
            if isinstance(src, dict):
                # support {"source": url, "from_url": True}
                url = src.get("source")
                fr = src.get("from_url", True)
            else:
                url = src
                fr = from_url
            if not url or not is_url(url):
                results.append({"source": url, "status": "error", "error_msg": "invalid url", "raw_text": "", "chunks": [], "embeddings": None})
                continue
            processed = process_document_for_llm(url, from_url=fr, do_embed=do_embed, chunk_size_chars=chunk_size_chars, overlap=overlap)
            results.append(processed)
        except Exception as e:
            results.append({"source": src, "status": "error", "error_msg": str(e), "raw_text": "", "chunks": [], "embeddings": None})
    return results

# -------------------------
# Retrieval utils
# -------------------------
def cosine_sim(a: Optional[np.ndarray], b: Optional[np.ndarray]) -> float:
    if a is None or b is None:
        return 0.0
    try:
        na = np.linalg.norm(a)
        nb = np.linalg.norm(b)
        if na == 0 or nb == 0:
            return 0.0
        return float(np.dot(a, b) / (na * nb))
    except Exception:
        return 0.0

def retrieve_top_k(query: str, processed_doc: Dict, top_k: int = 3, embed_model: str = "text-embedding-3-small"):
    chunks = processed_doc.get("chunks", []) or []
    embs = processed_doc.get("embeddings")
    if embs is not None and _client is not None:
        try:
            q_resp = _client.embeddings.create(model=embed_model, input=[query])
            q_emb = np.array(q_resp.data[0].embedding, dtype="float32")
            scores = [(i, float(cosine_sim(q_emb, np.array(e, dtype="float32")))) for i, e in enumerate(embs)]
            scores.sort(key=lambda x: x[1], reverse=True)
            results = []
            for idx, score in scores[:top_k]:
                ch = chunks[idx]
                results.append({"chunk_id": ch["id"], "text": ch["text"], "score": score})
            return results
        except Exception as e:
            LOGGER.warning("retrieve_top_k embedding path failed: %s", e)
    # substring fallback
    q = query.lower()
    matches = []
    for c in chunks:
        if q in c["text"].lower():
            matches.append({"chunk_id": c["id"], "text": c["text"], "score": 1.0})
        if len(matches) >= top_k:
            break
    if not matches:
        for c in chunks[:top_k]:
            matches.append({"chunk_id": c["id"], "text": c["text"], "score": 0.0})
    return matches
